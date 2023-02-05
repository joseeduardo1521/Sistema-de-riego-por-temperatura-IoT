from docx import Document

document = Document()

name = 'Eduardo'
phone = '3321296749'
email = 'josee375@outlook.com'

nm = 'Nombre'
ns = 'Telefono'
s  = 'Correo'
document.add_paragraph(nm + '|' + ns + '| ' + s)
document.add_paragraph(name + '|' + phone + '| ' + email)

document.save('test.docx')